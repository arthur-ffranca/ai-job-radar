from __future__ import annotations

import io
import re
import sys
import zipfile

from docx import Document
from pypdf import PdfReader

from app.services.profile_parser import EMPTY_PROFILE, parse_profile_with_debug


def _safe_debug_print(label: str, value: object = "") -> None:
    text = f"{label}{value}"
    encoding = sys.stdout.encoding or "utf-8"
    print(text.encode(encoding, errors="replace").decode(encoding, errors="replace"))


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf(contents: bytes) -> str:
    reader = PdfReader(io.BytesIO(contents))
    pages = [page.extract_text() or "" for page in reader.pages]
    return _clean_text("\n".join(pages))


def _extract_docx(contents: bytes) -> str:
    try:
        document = Document(io.BytesIO(contents))
    except zipfile.BadZipFile as exc:
        raise ValueError("DOCX text extraction failed. The file may be corrupted.") from exc

    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    ]
    return _clean_text("\n".join(paragraphs + table_cells))


def parse_resume_upload(
    filename: str,
    contents: bytes,
    content_type: str = "",
) -> dict[str, object]:
    lower = filename.lower()

    if lower.endswith(".pdf"):
        extraction_method = "pdf"
        raw_text = _extract_pdf(contents)
    elif lower.endswith(".docx"):
        extraction_method = "docx"
        raw_text = _extract_docx(contents)
    else:
        raise ValueError("Unsupported file type. Upload a PDF or DOCX resume.")

    _safe_debug_print("EXTRACTION METHOD: ", extraction_method)
    _safe_debug_print("RAW TEXT LENGTH: ", len(raw_text))
    _safe_debug_print("RAW TEXT PREVIEW:\n", raw_text[:1000])

    if not raw_text:
        raise ValueError("Text extraction returned empty content.")

    debug_info: dict[str, object] = {
        "filename": filename,
        "file_size": len(contents),
        "content_type": content_type,
        "extraction_method": extraction_method,
        "raw_text_length": len(raw_text),
        "raw_text_preview": raw_text[:1000],
    }

    if len(raw_text) < 300:
        error = "Resume text extraction failed or returned too little text."
        print("PARSER ERROR:", error)
        return {
            "error": error,
            "file_name": filename,
            "raw_text": raw_text,
            "profile": EMPTY_PROFILE,
            "raw_text_length": len(raw_text),
            "raw_text_preview": raw_text[:1000],
            "llm_response_preview": "",
            "parsed_profile": EMPTY_PROFILE,
            "debug": debug_info,
        }

    profile, parser_debug = parse_profile_with_debug(raw_text)
    debug_info.update(parser_debug)

    return {
        "file_name": filename,
        "raw_text": raw_text,
        "profile": profile,
        "raw_text_length": len(raw_text),
        "raw_text_preview": raw_text[:1000],
        "llm_response_preview": str(parser_debug.get("llm_response_preview", "")),
        "parsed_profile": profile,
        "debug": debug_info,
    }
